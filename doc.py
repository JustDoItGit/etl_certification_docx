# coding:utf-8
import sys
from docx import Document

# 此方法会出现重复，不建议使用
file = '../file_position/VCS-02590-A010-006_Rev3 签字终版.docx'
document = Document(file)
tables = document.tables

'''
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)
'''


# 将内容放到到三维列表里面
def make_list_3dim(tree):
    list_3dim = []
    t = 0
    for table in tree:
        list_3dim.append([])
        tr = 0
        for row in table.rows:
            list_3dim[t].append([])
            td = 0
            for cell in row.cells:
                text = (''.join(node.text for node in cell.paragraphs))
                list_3dim[t][tr].append(text.strip())
                print('t=%d,tr=%d,td=%d' % (t, tr, td))
                print(text)
                td += 1
            tr += 1
        t += 1
    return list_3dim


make_list_3dim(tables)
